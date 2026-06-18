from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)

NAVY    = RGBColor(0x1A, 0x25, 0x35)
SLATE   = RGBColor(0x3A, 0x5A, 0x7A)
GRANITE = RGBColor(0x8A, 0x88, 0x80)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
DARK    = RGBColor(0x22, 0x22, 0x22)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def h1(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(5)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '4');    bot.set(qn('w:color'), '1A2535')
    pBdr.append(bot); pPr.append(pBdr)

def h2(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = SLATE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)

def body(text, italic=False, color=None, sa=4):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = italic
    run.font.color.rgb = color if color else DARK
    p.paragraph_format.space_after = Pt(sa)

def note(text):
    body(text, italic=True, color=GRANITE, sa=4)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK

def step(num, title, detail=''):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{num}.  {title}')
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = NAVY
    if detail:
        r2 = p.add_run(f' — {detail}')
        r2.font.size = Pt(10); r2.font.color.rgb = DARK

def tbl(headers, rows, col_widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, row in enumerate(t.rows):
        for ci in range(len(headers)):
            row.cells[ci].width = Inches(col_widths[ci])
    hdr = t.rows[0]
    for ci, h in enumerate(headers):
        c = hdr.cells[ci]
        set_cell_bg(c, '1A2535'); set_cell_margins(c)
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
    for ri, rd in enumerate(rows):
        row = t.rows[ri+1]
        bg = 'F5F3F0' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(rd):
            c = row.cells[ci]
            set_cell_bg(c, bg); set_cell_margins(c)
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9); r.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── TITLE ──
p = doc.add_paragraph()
r = p.add_run('WAVEFORM MARKETING')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
r2 = p2.add_run('Services & Strategy — Working Document')
r2.font.size = Pt(13); r2.font.color.rgb = SLATE

p3 = doc.add_paragraph()
r3 = p3.add_run('June 2026  |  Confidential  |  Kim Nelson')
r3.font.size = Pt(10); r3.font.color.rgb = GRANITE
p3.paragraph_format.space_after = Pt(14)

# ── 1. PRICING PHILOSOPHY ──
h1('1. Pricing Philosophy')
body('Kim targets boutique fitness studio owners on the South Shore of MA. Competitors are primarily studio owners doing marketing themselves — not agencies. The value proposition is replacing poor DIY execution with 10+ years of professional growth marketing, website management, and marketing operations experience.', sa=5)
h2('Financial Targets')
bullet('Floor while at corporate: $2,000/mo')
bullet('Exit corporate threshold: $5,000/mo')
bullet('Ultimate target: 5 clients at $2,500/mo = $150K/year')
h2('Pricing Phases')
body('Phase 1 — Intro rates (first 2-3 paying clients): Lower end of ranges while building portfolio.', sa=2)
body('Phase 2 — Standard rates (after documented results): Full rates after 2-3 clients with measurable outcomes. Existing clients get 60 days notice.', sa=4)
h2('Free Work Policy')
body('1-2 website builds free to build portfolio. Cap at two. Require written testimonial, before/after metrics, and a signed agreement allowing case study use.', sa=2)
note('Important: Be Well Studios and similar ICP-quality prospects are NOT free clients. Save free builds for lower-stakes relationships.')

# ── 2. SERVICE MENU ──
h1('2. Full Service Menu')
tbl(
    ['Service','Phase 1','Phase 2','Type'],
    [
        ['Digital Presence Audit','$750','$1,000-$1,250','One-time'],
        ['Brand & Messaging Foundation','$1,500','$1,500','One-time'],
        ['Website Build (up to 8 pages)','$3,500','$5,000','One-time'],
        ['Additional pages (beyond 8)','$300/page','$300/page','Add-on'],
        ['Email Marketing Setup','$1,800','$1,800','One-time'],
        ['Email Template Package (3 templates)','$600','$600','One-time'],
        ['Website Care Plan','$500/mo','$600/mo','Retainer'],
        ['Marketing Operations Partner','$1,800/mo','$2,200-$2,500/mo','Retainer'],
        ['Full Partner','$2,800/mo','$3,000-$3,500/mo','Retainer'],
        ['Email Management add-on','+$400/mo','+$400/mo','Retainer add-on'],
    ],
    [3.0,1.3,1.5,1.2]
)

# ── 3. SERVICE DEFINITIONS ──
h1('3. Service Definitions')

h2('Digital Presence Audit — $750')
body('Website, Google Business Profile, SEO/AEO metadata, booking/lead flow, and existing email setup reviewed. Delivered as a written PDF report with prioritized findings and recommended next steps within 5 business days of intake form completion.', sa=3)
bullet('Not included: social media audit, paid advertising, print')
note('The audit deliverable IS the proposal. The last section includes a proposed site map, timeline, and investment summary. The contract and 50% deposit conversation happens at audit delivery.')

h2('Brand & Messaging Foundation — $1,500')
body('For studios that need to define what they are saying before a website can be built. Includes: brand positioning, core messaging (tagline, value proposition, service descriptions), about us narrative, content hierarchy, and a page-by-page copy brief.', sa=3)
note('Not included: logo design, visual identity, photography. Kim refers to a graphic designer in her network — client hires the designer directly.')

h2('Website Build — $3,500 (Phase 1) / $5,000 (Phase 2)')
body('Up to 8 pages. Additional pages at $300/page. Complex integrations quoted separately.', sa=3)
for item in [
    'Built in client\'s existing CMS where possible; Squarespace for net-new clients',
    'Brand-aligned design using client-provided assets',
    'Booking/registration integration, contact forms, lead capture',
    'On-page SEO/AEO setup (titles, meta, image alt text, sitemap, schema)',
    'Custom elements via Claude Code HTML export where needed',
    '2 rounds of revisions',
    'Full team training + written documentation handoff',
]:
    bullet(item)
body('Client provides: all copy, logo files (SVG + PNG), brand colors (hex), fonts, photography.', italic=True, color=GRANITE, sa=2)
body('Payment: 50% deposit at kickoff, 50% at launch.', sa=2)
note('Kill fee: If project is paused 30+ days due to client inaction, it is considered abandoned. Deposit is non-refundable. Restarting requires a new deposit.')

h2('Email Marketing Setup — $1,800 (add-on)')
body('For new builds: decided in discovery call, scoped alongside the website, built in from day one. For existing sites: standalone add-on available anytime.', sa=3)
for item in [
    'Website forms connected to email platform',
    'List and segment structure (lightweight CRM - e.g., Pilates list, CrossFit list)',
    '3 branded templates: newsletter, promotional, general',
    'Welcome sequence built and automated',
    'Full team training on how to clone, build, send, and manage lists',
]:
    bullet(item)
note('Platform: Kim works within client\'s existing platform. Brevo is the default for new clients.')

h2('Email Template Package — $600')
body('For clients who have email infrastructure but need fresh templates. Includes 3 branded templates (newsletter, promotional, general) + training. Additional templates beyond 3 priced per project.', sa=5)

h2('Website Care Plan — $500/mo (Phase 1) / $600/mo (Phase 2)')
body('Scope-based, not hours-based. Internal guideline: if a client consistently requires more than ~4 hours/month, the upgrade conversation to Ops Partner begins.', sa=3)
for item in [
    'Minor content updates: text, images, class schedules, link fixes',
    'Uptime monitoring',
    'Monthly SEO/AEO performance report (PDF)',
    'Employee onboarding to website (launch session only)',
]:
    bullet(item)
note('Not included: new pages or sections, structural design changes, new features, email management. Upgrade trigger: client regularly needs new sections or pages -> Ops Partner conversation.')

h2('Marketing Operations Partner — $1,800/mo (Phase 1) / $2,200-$2,500/mo (Phase 2)')
body('Includes everything in Care Plan, plus:', sa=2)
for item in [
    'New sections within existing pages',
    'Promotional and seasonal page updates',
    'Monthly 30-minute strategy call (Kim drives with prepared report)',
    'Monthly SEO/AEO monitoring and reporting (PDF)',
    'Lead workflow setup and optimization: CTA definition, sign-up flow, post-lead experience',
    'Ongoing employee onboarding to website',
    'Email template assistance: Kim guides, client executes',
    'Welcome sequence updates: Kim assists and guides',
    'Email Management add-on available: +$400/mo',
]:
    bullet(item)
note('Not included: new pages, new features, email campaign management (without add-on), social, paid ads, copywriting, photography, print. Upgrade trigger: client needs new pages, blog, or wants Kim to own email -> Full Partner conversation.')

h2('Full Partner — $2,800/mo (Phase 1) / $3,000-$3,500/mo (Phase 2)')
body('Includes everything in Ops Partner, plus:', sa=2)
for item in [
    'New pages built',
    'Blog management: topic suggestions, publishing, SEO/AEO optimization per post',
    'Writing assistance (Kim assists with structure; client provides final voice and approval)',
    'Email campaign management: builds, schedules, sends (requires Email Marketing Setup first)',
    'Welcome sequence updates: Kim executes',
    'Automation maintenance and updates',
    'Bi-monthly strategy sessions',
    'Deeper monthly reporting: website + email + lead performance combined',
    'Ongoing employee onboarding to website',
]:
    bullet(item)
note('Not included: social media management, paid ads, copywriting, photography, print.')

# ── 4. COMPARISON CHART ──
h1('4. Retainer Comparison Chart')
tbl(
    ['Feature','Care Plan\n$500/mo','Ops Partner\n$1,800/mo','Full Partner\n$2,800/mo'],
    [
        ['Text, image, schedule updates','Yes','Yes','Yes'],
        ['Link fixes, form updates','Yes','Yes','Yes'],
        ['New sections within existing pages','No','Yes','Yes'],
        ['Promotional/seasonal page updates','No','Yes','Yes'],
        ['New pages','No','No','Yes'],
        ['Blog publishing + SEO/AEO','No','No','Yes'],
        ['New features or integrations','No','No','Yes'],
        ['Employee onboarding to website','Launch only','Ongoing','Ongoing'],
        ['Monthly 30-min strategy call','No','Yes','Yes'],
        ['Lead workflow setup + optimization','No','Yes','Yes'],
        ['Monthly SEO/AEO report (PDF)','Yes','Yes','Yes'],
        ['Email template changes','Not included','Kim assists','Kim does it'],
        ['Welcome sequence updates','Not included','Kim assists','Kim does it'],
        ['Email newsletter management','+$400/mo','+$400/mo','Included'],
        ['Blog strategy + writing assist','No','No','Yes'],
    ],
    [3.2,1.4,1.4,1.4]
)

# ── 5. CLIENT PROFILES ──
h1('5. Client Profiles')
h2('Care Plan Client')
body('Has a working website. Needs it maintained and reported on. Low-touch, predictable. Typically a post-build client who does not need ongoing strategic input.', sa=5)
h2('Marketing Operations Partner Client')
body('Foundation is in place. Needs someone to own the infrastructure, optimize lead flow, and bring strategic thinking monthly. Can execute on their own when given direction. Example: Be Well Studios.', sa=5)
h2('Full Partner Client')
body('Wants Kim embedded. Busier, more complex, or growing fast. Wants content strategy, blog management, and email handled. May have multiple service lines or an owner without bandwidth for content. Buying a marketing brain, not just maintenance.', sa=3)
note('Shorthand: Ops Partner clients execute. Full Partner clients hand it off.')

# ── 6. EMAIL PRICING LADDER ──
h1('6. Email Marketing Pricing Ladder')
tbl(
    ['Service','Price','When to use'],
    [
        ['Email Template Package','$600','Client has platform in place, needs templates refreshed'],
        ['Email Marketing Setup','$1,800','Client needs full infrastructure built from scratch'],
        ['Email Management add-on','+$400/mo','Client wants Kim to manage and send campaigns'],
        ['Full Partner','Included','Full email management included in retainer'],
    ],
    [2.2,1.2,4.0]
)
body('How email integrates with the website: Website forms -> email platform lists -> list/segment structure acts as lightweight CRM. Set up once during Email Marketing Setup, trained on at handoff.', italic=True, color=GRANITE, sa=5)

# ── 7. WEBSITE BUILD PROCESS ──
h1('7. Website Build Process (14 Steps)')
build_steps = [
    ('Intake form','Client completes before any call.'),
    ('Discovery call','60 min video. Kim drives agenda. Granola captures notes.'),
    ('Scope document','Written. Client signs off before anything starts.'),
    ('50% deposit collected','Timeline does not start until payment clears.'),
    ('Full access checklist collected','See Section 8. Kim helps set up anything the client does not have.'),
    ('Screaming Frog crawl + landscape review','Full picture of existing site before any concepts are presented.'),
    ('Concept deck','Built in Lovable from discovery notes + crawl findings. Presented for approval before build begins.'),
    ('Build','Starts after deck is approved.'),
    ('Round 1 review','Staging link shared. All feedback batched in one round.'),
    ('Round 2 review','Final changes only. No new scope.'),
    ('QA + Launch',''),
    ('50% balance collected at launch',''),
    ('Training + documentation handoff','Full team trained. Written docs delivered.'),
    ('Retainer conversation begins',''),
]
for i, (t, d) in enumerate(build_steps, 1):
    step(i, t, d)
note('\nKill fee: If project is paused 30+ days due to client inaction, it is considered abandoned. Deposit is non-refundable. Restarting requires a new deposit.')

# ── 8. ACCESS CHECKLIST ──
h1('8. Client Access Checklist')
body('Collected in full at intake. If a client does not have something, Kim either skips it or sets it up. Nothing starts until this checklist is complete.', sa=5)
h2('Brand Assets (client provides)')
for item in ['Logo files: SVG and PNG minimum','Brand colors: hex codes','Brand fonts or font preferences','Photography and images','All website copy (written and approved)','Any existing content to carry over']:
    bullet(item)
h2('Technical Access (never via email or text)')
for item in [
    'Domain registrar login — to update DNS records (different from website builder login)',
    'Existing website CMS login — for content reference and migration',
    'Google Analytics 4 — admin preferred, read acceptable',
    'Google Search Console — admin preferred; Kim sets up if they do not have it',
    'Google Business Profile — admin access',
    'Google Tag Manager — if exists',
    'Booking software admin: Mindbody, Vagaro, Acuity, Pike13, etc.',
    'Email marketing platform — if exists',
]:
    bullet(item)
note('Security rule: Credentials are never shared over email or text. A password manager or secure sharing tool is used for all transfers.')

# ── 9. AUDIT PROCESS ──
h1('9. Audit Process (6 Steps)')
audit_steps = [
    ('Discovery call','Goals, brand vision, how visitors should feel, what success looks like. Must happen before any audit work begins.'),
    ('Access collection','GA4 (admin preferred) + Search Console (Kim sets up if not present) + current website login.'),
    ('Screaming Frog crawl','Full site inventory: pages, response codes, broken links, page structure. Combined with GA4 shows what exists and how it is performing.'),
    ('Compile the report','Repeatable template, simple language, sectioned by category. Each finding: what is happening -> what I would do -> what it means for your business.'),
    ('Deliver','PDF sent within 5 business days. Present on video or in person if the vibe is right.'),
    ('Contract + 50% deposit','Audit delivery IS the close. Report ends with proposed site map, timeline, and investment. Contract conversation flows from the findings.'),
]
for i, (t, d) in enumerate(audit_steps, 1):
    step(i, t, d)

# ── 10. AUDIT REPORT TEMPLATE ──
h1('10. Audit Report Template Structure (8-12 pages)')
report_secs = [
    ('Cover page','Client name, date, Waveform Marketing. One sentence capturing the single biggest opportunity.'),
    ('Section 1 - What We Reviewed','Tools used (Screaming Frog, GA4, Search Console, manual review). Establishes credibility.'),
    ('Section 2 - Your Current Website','Plain-language summary: pages, what is working, what is broken, what is missing. No jargon.'),
    ('Section 3 - How You Show Up Online','SEO/AEO health from the customer perspective. Keywords, rankings, metadata, GBP, Search Console, AI search visibility.'),
    ('Section 4 - How You Capture Leads','CTAs, booking flow, forms, post-submission experience. Walk the full lead journey from first click to member.'),
    ('Section 5 - Email & Follow-Up','List status, automations, post-sign-up experience. Gaps flagged with Email Marketing Setup as solution.'),
    ('Section 6 - Priority Findings','High / Medium / Low priority. Each finding: what is happening -> what I would do -> what it means for your business.'),
    ('Section 7 - What I Would Build','Proposed site map. Pages to migrate, cut, add - with one-line rationale tied to discovery call goals.'),
    ('Section 8 - What This Means for Your Business','Outcomes tied to what the client said they wanted. Concrete and business-specific, not generic SEO language.'),
    ('Section 9 - Timeline & Investment','Visual project timeline scaled to site complexity. Investment summary. Clear next step: sign, deposit, set kickoff date.'),
]
for title, detail in report_secs:
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title)
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = SLATE
    r2 = p.add_run(f'\n{detail}')
    r2.font.size = Pt(10); r2.font.color.rgb = DARK

# ── 11. TECH STACK ──
h1('11. Tech Stack')
tbl(
    ['Layer','Tool','Notes'],
    [
        ['CMS (net-new clients)','Squarespace','Default. Business plan minimum required for custom code injection.'],
        ['CMS (existing clients)','Client\'s current platform','Stay where they are: WordPress, Wix, Squarespace'],
        ['CMS (unworkable)','Migrate to Squarespace','Custom-built with no CMS, or too outdated'],
        ['Custom elements','Claude Code -> HTML embed','Bespoke components embedded into any CMS'],
        ['Email marketing','Client\'s existing platform','Brevo as default for new/no-platform clients'],
        ['Booking integrations','Client\'s existing platform','Mindbody, Vagaro, Acuity, Pike13 - embedded'],
        ['Analytics','GA4 + Search Console','Client owns accounts; Kim has admin access'],
        ['Forms + lead capture','Platform-native','Connected to email platform'],
        ['SEO/AEO audit tool','Screaming Frog','Kim\'s tool; not client-facing'],
    ],
    [1.8,1.8,3.8]
)
h2('CMS Decision Logic')
tbl(
    ['Situation','Platform'],
    [
        ['Client has no website','Squarespace (default)'],
        ['Client is on WordPress, Wix, or Squarespace','Stay on their platform'],
        ['Client is on something custom or unworkable','Migration conversation -> Squarespace'],
    ],
    [4.2,3.2]
)
h2('Licensing Rule')
body('Client always owns and pays for their own accounts. Kim is added as admin contributor. If the relationship ends, everything stays with the client and Kim removes herself. Kim never pays for client subscriptions.', sa=5)

# ── 12. OUT OF SCOPE ──
h1('12. What\'s Out of Scope')
for item in [
    'Social media management (by design: "I work alongside your social person and handle everything they can\'t")',
    'Paid advertising (Google Ads, Meta Ads, etc.)',
    'Print marketing',
    'Photography',
    'Video editing',
    'Logo and visual identity design (refer to network designer)',
]:
    bullet(item)
note('Out-of-scope requests: Kim refers clients to agency partners. Referral relationships to formalize: (1) graphic designer for logo/visual identity, (2) social media agency serving fitness studios.')

# ── 13. WORKLOAD ANALYSIS ──
h1('13. Workload Analysis - 3 Clients')
tbl(
    ['Retainer','Hrs/Month','Hrs/Week','Effective Rate'],
    [
        ['Website Care Plan','3-5 hrs','<1 hr','~$100-165/hr'],
        ['Marketing Operations Partner','9-13 hrs','2-3 hrs','~$138-200/hr'],
        ['Full Partner','14-20 hrs','3-5 hrs','~$140-200/hr'],
    ],
    [2.8,1.4,1.4,1.8]
)
h2('Scenario A: 2 Ops Partner + 1 Care Plan')
bullet('~26 hours/month | ~6-7 hours/week'); bullet('Revenue: $4,100/mo')
h2('Scenario B: 1 Full Partner + 1 Ops Partner + 1 Care Plan')
bullet('~32 hours/month | ~8 hours/week'); bullet('Revenue: $5,100/mo')
note('3 clients at target revenue is roughly one full work day per week - fits inside a 9-2 schedule. Heavier lift is upfront (builds, email setup). Once on retainer, very manageable alongside a corporate job.')

# ── 14. CLIENT JOURNEYS ──
h1('14. Client Journey Examples')
tbl(
    ['Journey Type','Services','Project Revenue'],
    [
        ['Client with nothing','Brand & Messaging ($1,500) -> Designer referral -> Website Build ($3,500) -> Email Setup ($1,800) -> Retainer','$6,800'],
        ['Standard client','Audit ($750) -> Website Build + Email Setup ($5,300) -> Retainer','$6,050'],
        ['Audit-to-build','Audit ($750) -> Contract signed -> Website Build ($3,500) -> Retainer','$4,250'],
        ['Be Well Studios (est.)','Audit ($750) -> Build + Email Setup ($5,300) -> $1,800/mo retainer','$28,650 yr 1'],
    ],
    [1.8,4.0,1.5]
)

# ── 15. STILL TO DEFINE ──
h1('15. Still to Define')
h2('Business Operations')
for item in [
    'Retainer contract terms: notice period, minimum commitment length, pause policy',
    'Proposal structure: what Kim actually sends a prospect to close a deal',
    'Phase 2 pricing triggers: specific milestone that moves from intro to standard rates',
    'Referral partner formalization: graphic designer + social media agency',
]:
    bullet(item)
h2('Deliverables')
bullet('Reporting templates: what the monthly SEO/AEO report actually looks like')
h2('From Original Strategy Document')
for item in ['LLC formation timing','Corporate exit financial runway','Consulting schedule structure alongside dance classes']:
    bullet(item)

# ── FOOTER ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
r = p.add_run('Waveform Marketing  |  kim@waveformmarketing.com  |  waveformmarketing.com  |  Confidential - June 2026')
r.font.size = Pt(8); r.font.color.rgb = GRANITE

out = '/sessions/sharp-great-galileo/mnt/outputs/Waveform_Marketing_Strategy.docx'
doc.save(out)
print(f'Saved: {out}')
